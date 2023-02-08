import pyttsx3
import datetime
import speech_recognition as sr
import wikipedia
import webbrowser;' '
import os
import random
from email.message import  EmailMessage
import ssl
import smtplib
import pywhatkit as kit 
import webbrowser
import sys 
import requests
from bs4 import BeautifulSoup
from random import randrange
from turtle import *
from freegames import square, vector
from freegames import line
from playsound import playsound
import pyjokes
from docx import Document
import random 
import pyautogui
import wave
import pyaudio
from ctypes import cast,POINTER
from comtypes import CLSCTX_ALL
from pycaw.pycaw import AudioUtilities,IAudioEndpointVolume
from pywhatkit.core.exceptions import InternetException
import pathlib
import time
from platform import system
from urllib.parse import quote
from webbrowser import open









engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
# to give the voice code pipiprint(voices[1].id)
engine.setProperty('voice', voices[1].id)

def speak(audio):
    engine.say(audio)
    engine.runAndWait()

def takecommand():
    #It takes microphone input        
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        r.pause_threshold = 1
        audio=r.listen(source)

    try:
        print("Recognizing...")
        query = r.recognize_google(audio, language='en-in')
        print(f"User said: {query}\n")
        #speak(query)

    except Exception as e:
        # it shows error print(e)
        print("Say that again please")
        speak("Say that again please")
        return "None"
    return query 

def wishme():
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("good morning.") 

    elif hour>=12 and hour<18:
        speak("Good Afternoon.")

    else:
        speak("Good Evening.")

    speak("Hello!! I am Friday ,how may i help you today")



#email Sender

def emailtoharsh():
    email_sender='hv5699012'
    email_password= 'qxxgfqeyghgisqdo'
    email_receiver='harshvermajpr1@gmail.com'
    speak("Tell me the subject")
    subject=takecommand()
    print("Subject of the email:- ",subject)
    speak("Tell me the body of the email.")
    body=takecommand()
    print("Body of the email:- ",body)
    em=EmailMessage()
    em['From']=email_sender
    em['To']=email_receiver
    em['Subject']=subject
    em.set_content(body)

    context=ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com',465,context=context) as smtp:
        smtp.login(email_sender,email_password)
        smtp.sendmail(email_sender,email_receiver,em.as_string())

    print("Email Sent.")
    speak("Email Sent.")


def emailtokhushi():
    email_sender='hv5699012'
    email_password= 'qxxgfqeyghgisqdo'
    email_receiver='ratrakr@gmail.com'
    speak("Tell me the subject")
    subject=takecommand()
    print("Subject of the email:- ",subject)
    speak("Tell me the body of the email.")
    body=takecommand()
    print("Body of the email:- ",body)
    em=EmailMessage()
    em['From']=email_sender
    em['To']=email_receiver
    em['Subject']=subject
    em.set_content(body)

    context=ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com',465,context=context) as smtp:
        smtp.login(email_sender,email_password)
        smtp.sendmail(email_sender,email_receiver,em.as_string())

    print("Email Sent.")
    speak("Email Sent.")


#Play games

#1. play snake game 
def snakegame():
    food = vector(0, 0)
    snake = [vector(10, 0)]
    aim = vector(0, -10)


    def change(x, y):
        """Change snake direction."""
        aim.x = x
        aim.y = y


    def inside(head):
        """Return True if head inside boundaries."""
        return -200 < head.x < 190 and -200 < head.y < 190


    def move():
        """Move snake forward one segment."""
        head = snake[-1].copy()
        head.move(aim)

        if not inside(head) or head in snake:
            square(head.x, head.y, 9, 'red')
            update()
            return

        snake.append(head)

        if head == food:
            print('Snake:', len(snake))
            food.x = randrange(-15, 15) * 10
            food.y = randrange(-15, 15) * 10
        else:
            snake.pop(0)

        clear()

        for body in snake:
            square(body.x, body.y, 9, 'black')

        square(food.x, food.y, 9, 'green')
        update()
        ontimer(move, 100)


    setup(420, 420, 370, 0)
    hideturtle()
    tracer(False)
    listen()
    onkey(lambda: change(10, 0), 'Right')
    onkey(lambda: change(-10, 0), 'Left')
    onkey(lambda: change(0, 10), 'Up')
    onkey(lambda: change(0, -10), 'Down')
    move()
    done()

def tictactoe():
    
    def grid():
        """Draw tic-tac-toe grid."""
        line(-67, 200, -67, -200)
        line(67, 200, 67, -200)
        line(-200, -67, 200, -67)
        line(-200, 67, 200, 67)


    def drawx(x, y):
        """Draw X player."""
        line(x, y, x + 133, y + 133)
        line(x, y + 133, x + 133, y)


    def drawo(x, y):
        """Draw O player."""
        up()
        goto(x + 67, y + 5)
        down()
        circle(62)


    def floor(value):
        """Round value down to grid with square size 133."""
        return ((value + 200) // 133) * 133 - 200


    state = {'player': 0}
    players = [drawx, drawo]


    def tap(x, y):
        """Draw X or O in tapped square."""
        x = floor(x)
        y = floor(y)
        player = state['player']
        draw = players[player]
        draw(x, y)
        update()
        state['player'] = not player


    setup(420, 420, 370, 0)
    hideturtle()
    tracer(False)
    grid()
    update()
    onscreenclick(tap)
    done()

#Set an alarm

"""def setalarm():
    os. system('clear')
    alarmH = int(input("What hour do you want the alarm to ring? "))
    alarmM = int(input("What minute do you want the alarm to ring? "))
    amPm = str(input("am or pm? "))
    os. system('clear')
    print("Waiting for alarm",alarmH,alarmM,amPm)
    if (amPm == "pm"):
            alarmH = alarmH + 12
    while(1 == 1):
        if(alarmH == datetime.datetime.now().hour and
            alarmM == datetime.datetime.now().minute) :
            print("Time to wake up")
            playsound('/your/path/to/file/beep-07.mp3')
            break"""

#Write a word file.
def wordfile():
    document=Document()
    print("Tell me the Heading.")
    speak("Tell me the Heading.")
                    
    head=takecommand()
    document.add_heading(head, 0)
    print("Tell me the Body.")
    speak("Tell me the Body.")
                    
    body=takecommand()
    document.add_paragraph(body).bold= True
    document.save("C:/Users/Harsh Verma/Desktop/Friday/Notepad file accessing/Document.docx")
    speak("done writing word file")

#Random password generator
def genpass():
    uppercase_letters="ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    lowercase_letters=uppercase_letters.lower()
    digits="1234567890"
    upper,lower,nums=True,True,True 
    al=""
    if upper:
        al+=uppercase_letters

    if lower:
        al+=lowercase_letters

    if digits:
        al+=digits

    length=8
    amount=1

    for x in range(amount):
        password="".join(random.sample(al,length))
        
        print("Password Generated: {}".format(password))
        speak("Password Generated.")
        print("")

def screeninshot():
    screenshot=pyautogui.screenshot()
    uppercase_letters="ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    lowercase_letters=uppercase_letters.lower()
    digits="1234567890"
    upper,lower,nums=True,True,True 
    al=""
    if upper:
        al+=uppercase_letters

    if lower:
        al+=lowercase_letters

    if digits:
        al+=digits

    length=8
    amount=1

    for x in range(amount):
        password="".join(random.sample(al,length))
        a=".png"
        new=(password+a)
        print("Image name: {}".format(new))
        screenshot.save("C://Users//Harsh Verma//Desktop//Friday//Notepad file accessing//{}".format(new)) 
        speak("Screenshot Taken.")
        print("Screenshot Taken.")   

#Voice Recorder
def voicerecorder():
    print("Recording start.")
    speak("Recording start.")
    
    audio=pyaudio.PyAudio()
    stream=audio.open(format=pyaudio.paInt16,channels=1,rate=44100,input=True,frames_per_buffer=1024)
    frames=[]
    try:
        while True:
            data=stream.read(1024)
            frames.append(data)

    except KeyboardInterrupt:
        pass

    stream.stop_stream()
    stream.close()
    audio.terminate()
    sound_file=wave.open("C://Users//Harsh Verma//Desktop//Friday//Notepad file accessing//recordedaudio.wav","wb")
    sound_file.setnchannels(1)
    sound_file.setsampwidth(audio.get_sample_size(pyaudio.paInt16))
    sound_file.setframerate(44100)
    sound_file.writeframes(b''.join(frames))
    sound_file.close()
    print("Voice Recorded.")
    speak("Voice Recorded.")

def volsetup():
    devices=AudioUtilities.GetSpeakers()    
    interface=devices.Activate(IAudioEndpointVolume._iid_,CLSCTX_ALL, None)
    volume=cast(interface,POINTER(IAudioEndpointVolume))
    if ' 0%' in query:
        volume.SetMasterVolumeLevel(-65.0, None)
        speak("Volume Level Changed to Zero.")
    elif ' 100%' in query:
        volume.SetMasterVolumeLevel(0.0, None)
        speak("Volume Level Changed to Full.")
    elif ' 50%' in query:
        volume.SetMasterVolumeLevel(-10.0, None)
        speak("Volume Level Changed to Half.")

    elif ' 30%' in query:
        volume.SetMasterVolumeLevel(-18.0, None)
        speak("Volume Level Changed to 30 Percent.")

    elif ' 70%' in query:
        volume.SetMasterVolumeLevel(-5.0, None)
        speak("Volume Level Changed to 70 Percent.")

def add():
    pass

#Emergency Calls and help.
def emergency():
    y=''
    while(y==''):
        print("Emergency Mode On.")
        speak("Emergency Mode On.")
        devices=AudioUtilities.GetSpeakers()    
        interface=devices.Activate(IAudioEndpointVolume._iid_,CLSCTX_ALL, None)
        volume=cast(interface,POINTER(IAudioEndpointVolume))
        volume.SetMasterVolumeLevel(0.0, None)

        #alert sound
        def alertsound():
            music_dir = "C:\\Users\\Harsh Verma\\Desktop\\Friday\\Popular Emergency Alarm Clock  Danger Alarm  Sound Effect.mp3"
            os.startfile(os.path.join(music_dir))
        
        alertsound()
        
        try:
            
            content= 'Help I Need Help as soon as possible.'
            hour = int(datetime.datetime.now().hour)
            minute = int(datetime.datetime.now().minute)
            phone='+91 63679 82717'
            kit.sendwhatmsg_instantly(phone, content , hour , minute)
            print("Message Sent sir.")
            speak("message Sent sir.")   

        except Exception as e:
            print(e)  
            speak("Sorry sir . i can't do that right now")
        
        alertsound()
        
        #email
        email_sender='hv5699012'
        email_password= 'qxxgfqeyghgisqdo'
        email_receiver='harshvermajpr1@gmail.com'
        subject=("Urgent Help Needed.")
        print("Subject of the email:- ",subject)
        body=("I am in an emergency. Please help me as soon as possible.")
        print("Body of the email:- ",body)
        em=EmailMessage()
        em['From']=email_sender
        em['To']=email_receiver
        em['Subject']=subject
        em.set_content(body)

        context=ssl.create_default_context()

        with smtplib.SMTP_SSL('smtp.gmail.com',465,context=context) as smtp:
            smtp.login(email_sender,email_password)
            smtp.sendmail(email_sender,email_receiver,em.as_string())

        print("Email Sent For your help.")
        speak("Email Sent For your help.")

        y=takecommand()
        if y=='stop':
            continue
        else:
            y=''

    


#Main Program
if __name__ == "__main__":
    wishme()
    while True:
#if 1:
        query = takecommand().lower()
#logical things
        if 'introduce yourself' in query:
            speak('Hey. i am friday.  i am basically an assistant , who can do some amazing things , like i can play music if you want , i can send emails , i can open folders you want , and many more .  you wanna try something ? ')
        
        elif 'sing' in query:
            str=speak('By the way im not too good at singing but i can try...')
            song=('Tumko paaya hai to jaise khoya hoon kehna chahoon bhi to tumse kya kahoo kisi zabaan mein bhivo lavz hi nahi ki jinme tum ho kya tumhe bata sakoon main agar kahoon tumsa haseen')
            song1=('If you wanna run away with me, I know a galaxy And I can take you for a ride......... I had a premonition that we fell into a rhythm Where the music dont stop for life....... Glitter in the sky.... glitter in my eyes..... Shining just the way I like. If youre feeling like you need a little bit of company You met me at the perfect time')
            speak(song1)
            speak("I can sing very well..right?")




        elif 'wikipedia' in query:
            speak("Searching Wikipedia...")
            query = query.replace("wikipedia", "")
            results = wikipedia.summary(query, sentences = 2)
            speak("According to Wikipedia")
            print(results)
            speak(results)
            
#WEB BASED FUNCTIONS.    
        elif 'open youtube' in query:
            speak("Opening Youtube...")
            webbrowser.open("https://www.youtube.com/")

        elif 'open google.com' in query:
            speak("Opening Google...")
            webbrowser.open("https://www.google.com/")

        elif 'open youtube trending page' in query:
            speak("Opening youtube trending page...")
            webbrowser.open(" https://www.youtube.com/feed/trending?bp=6gQJRkVleHBsb3Jl")

        elif 'open amazon.com' in query:
            speak("Opening Amazon")
            webbrowser.open("https://www.amazon.in/")

        elif 'open gmail.com' in query:
            speak("Opening Gmail")
            webbrowser.open("https://mail.google.com")

        elif 'open maps.com' in query:
            speak("Opening Google Maps")
            webbrowser.open("https://www.google.com/maps/@26.6307398,73.873897,7z")

        elif 'open translate.com' in query:
            speak("Opening Google Translate")
            webbrowser.open("https://translate.google.co.in/")

        elif 'open flipkart.com' in query:
            speak("Opening Flipkart")
            webbrowser.open("https://www.flipkart.com/")

        elif 'open github.com' in query:
            speak("Opening Github")
            webbrowser.open("https://github.com/")

        elif 'open googledrive.com' in query:
            speak("Opening Google drive")
            webbrowser.open("https://drive.google.com/drive/my-drive")

#Play Music
        elif 'open music' in query:
            n= random.randint(0,18)
            print(n)
            speak("Here WE Go")
            music_dir = 'C:\\Users\\Harsh Verma\\Music'
            songs= os.listdir(music_dir)
            speak("Playing Song Number")
            speak(n)
            print(songs)
            os.startfile(os.path.join(music_dir,songs[n]))

        elif 'change music' in query:
            n= random.randint(0,18)
            print(n)
            speak("Okay.")
            music_dir = 'C:\\Users\\Harsh Verma\\Music'
            songs= os.listdir(music_dir)
            speak("Playing Song Number")
            speak("Playing Song Number")
            speak(n)
            print(songs)
            os.startfile(os.path.join(music_dir,songs[n]))

#Software Opening 

        elif 'open vs code' in query:
            vscodepath="C:\\Users\\Harsh Verma\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe"
            os.startfile(vscodepath)
            speak("Visual Studio Code opened.")

        elif 'close vs code' in query:
            os.system("taskkill /f /im Code.exe")
            speak("Visual Studio Code has been Closed.")
            
        elif 'open visual studio' in query:
            vscodepath="C:\\Program Files\\Microsoft Visual Studio\\2022\\Community\\Common7\\IDE\\devenv.exe"
            os.startfile(vscodepath)
            speak("Visual Studio opened.")

        elif 'close visual studio' in query:
            os.system("taskkill /f /im devenv.exe")
            speak("Visual Studio has been Closed.")

        elif 'open android studio' in query:
            vscodepath="D:\\Android Studio\\bin\\studio64.exe"
            os.startfile(vscodepath)
            speak("Android Studio opened.")

        elif 'close android studio' in query:
            os.system("taskkill /f /im studio64.exe")
            speak("Android Studio has been Closed.")

        elif 'open brave' in query:
            vscodepath="C:\\Program Files\\BraveSoftware\\Brave-Browser\\Application\\brave.exe"
            os.startfile(vscodepath)
            speak("Brave opened.")

        elif 'close brave' in query:
            os.system("taskkill /f /im brave.exe")
            speak("Brave has been Closed.")

        elif 'open google chrome' in query:
            vscodepath="C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe"
            os.startfile(vscodepath)
            speak("Google Chrome opened.")

        elif 'close google chrome' in query:
            os.system("taskkill /f /im chrome.exe")
            speak("Google Chrome has been Closed.")

        elif 'open discord' in query:
            vscodepath="C:\\Users\\Harsh Verma\\AppData\\Local\\Discord\\Update.exe"
            os.startfile(vscodepath)
            speak("Discord opened.")

        elif 'close discord' in query:
            os.system("taskkill /f /im Discord.exe")
            speak("Discord has been Closed.")


#email sender

        elif 'send email to harsh' in query:
            emailtoharsh()

        elif 'send email to khushi' in query:
            emailtokhushi()

#play youtube videos          
        elif 'play' in query:
            query = query.replace("play", "")
            results = kit.playonyt(query)
            str=('playing',query)
            speak(str)
            print(results)

#Google Search Anything
        elif 'google' in query:
            query = query.replace("google", "")
            results = kit.search(query)
            str=('searching for',query)
            speak(str)
            print("Searching results for", query )

#send whatsapp messages


        elif 'send a message to' in query:
            query=query.replace("send a message to ","")
            try:
                speak('What should i say?')
                content= takecommand()
                hour = int(datetime.datetime.now().hour)
                minute = int(datetime.datetime.now().minute)
                if content=='None':
                    content=content.replace(content , "")

                if query=="pankaj":
                    phone='+91 70236 22705'

                elif query=="khushi":
                    phone='+91 88825 30012'

                elif query=="chandu":
                    phone='+91 80056 46802'

                elif query=="priya":
                    phone='+91 73398 98849'

                elif query=="itisha":
                    phone='+91 74269 32093'

                elif query=="roshan":
                    phone='+91 97841 17842'

                elif query=="harjeet sir":
                    phone='+91 81074 54956'


             
                else :
                    print("contact not found.")
                    speak("contact not found.")

                kit.sendwhatmsg_instantly(phone, content , hour , minute)
                print("Message Sent sir.")
                speak("message Sent sir.")
            except Exception as e:
                print(e)  
                speak("Sorry sir . i can't do that right now")

        elif 'thanks' in query:
            speak('no problem sir.')

        elif 'run snake game' in query:
            snakegame()

        elif 'close voice assistant' in query:
            speak("see you soon buddy.")
            break

        elif 'temperature in' in query:
            query=query.replace("temperature in","")
            print(query)
            print("Displaying Weather Report For :- "+query)
            speak("Displaying Weather Report For :- "+query)
            url='https://wttr.in/{}'.format(query)
            res = requests.get(url)
            print(res.text) 

        elif 'shutdown the pc' in query:
                speak("Okay. your PC is going to be shutdown in 2 minutes")
                kit.shutdown(time=120)
                
        elif 'cancel shutdown' in query:
                kit.cancel_shutdown()
                speak("Shutdown has been cancelled.")

            
        elif 'joke' in query:
            a=pyjokes.get_joke()
            print(a) 
            speak(a)

        elif 'spell' in query:
            len=len(query)
            speak(query[0],query[1])
        
        elif 'write a notepad' in query:
            query=query.replace("write a note","")
            speak("tell me what you want to write.")
            query=takecommand()
            f=open("C:/Users/Harsh Verma/Desktop/Friday/Notepad file accessing/new.txt","w")
            f.write(query)
            f.close()
            speak("File created.")

        elif 'write a word file' in query:
            wordfile()

        elif "generate password" in query:
            genpass()

        elif 'take screenshot' in query:
            screeninshot()

        elif 'start recording' in query:
            voicerecorder()

        elif 'set volume level' in query:
            query=query.replace("set volume level","")
            volsetup()

        elif 'mute' in query:
            devices=AudioUtilities.GetSpeakers()    
            interface=devices.Activate(IAudioEndpointVolume._iid_,CLSCTX_ALL, None)
            volume=cast(interface,POINTER(IAudioEndpointVolume))
            volume.SetMute(0,None)

        elif 'unmute' in query:
            devices=AudioUtilities.GetSpeakers()    
            interface=devices.Activate(IAudioEndpointVolume._iid_,CLSCTX_ALL, None)
            volume=cast(interface,POINTER(IAudioEndpointVolume))
            volume.SetMute(1,None)

        elif 'add' in query:
            pass


        elif 'emergency' in query:
            emergency()

        elif 'open whatsapp' in query:
            kit.open_web()

        


        

            





        



       